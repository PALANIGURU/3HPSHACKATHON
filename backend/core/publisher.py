"""
publisher.py — Step 8: Render a .docx shift handover report.

Rules:
- Always render all 4 sections (Completed / In Progress / Blockers / Watch-list).
- Empty sections show "Nothing to report." — never omit a section.
- Raises RuntimeError loudly if export fails (non-zero exit in CLI / 500 in API).
- Uses python-docx for document generation.
"""

import os
from datetime import datetime
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── Section config ─────────────────────────────────────────────────────────────

SECTION_CONFIG = [
    {
        "key": "completed",
        "title": "✅  Completed",
        "heading_color": RGBColor(0x1A, 0x7A, 0x45),  # Dark green
        "empty_msg": "Nothing to report.",
    },
    {
        "key": "in_progress",
        "title": "🔄  In Progress",
        "heading_color": RGBColor(0x1D, 0x4E, 0xD8),  # Blue
        "empty_msg": "Nothing to report.",
    },
    {
        "key": "blockers",
        "title": "🚨  Blockers",
        "heading_color": RGBColor(0xB9, 0x1C, 0x1C),  # Red
        "empty_msg": "Nothing to report.",
    },
    {
        "key": "watch_list",
        "title": "👁️  Watch-list",
        "heading_color": RGBColor(0xD9, 0x77, 0x06),  # Amber
        "empty_msg": "Nothing to report.",
    },
]


def _set_paragraph_spacing(para, space_before: int = 6, space_after: int = 3):
    """Set paragraph spacing in points."""
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)


def _add_horizontal_rule(doc: Document):
    """Add a thin horizontal rule paragraph."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after = Pt(2)
    # Add bottom border to simulate HR
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _format_event_line(item: dict) -> str:
    """Format a single event item as a bullet line."""
    return f"[{item['record_id']}] {item['summary']}  (status: {item['status']}, {item['timestamp']})"


def render_docx(
    sections: dict,
    output_path: str,
    shift_start: Optional[datetime] = None,
    shift_end: Optional[datetime] = None,
    generated_at: Optional[datetime] = None,
) -> str:
    """
    Render the shift handover report as a .docx file.

    Args:
        sections:     Output from generator.generate_sections() — 4-key dict.
        output_path:  Absolute path where the .docx should be saved.
        shift_start:  Optional shift start datetime for the report header.
        shift_end:    Optional shift end datetime for the report header.
        generated_at: Optional generation timestamp; defaults to now (UTC).

    Returns:
        The resolved output_path string.

    Raises:
        RuntimeError: If the document cannot be created or written.
    """
    try:
        if generated_at is None:
            from datetime import timezone
            generated_at = datetime.now(tz=timezone.utc)

        doc = Document()

        # ── Page margins ─────────────────────────────────────────────────────
        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # ── Document Title ────────────────────────────────────────────────────
        title_para = doc.add_paragraph()
        title_run = title_para.add_run("Shift Handover Report")
        title_run.bold = True
        title_run.font.size = Pt(20)
        title_run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(title_para, space_before=0, space_after=4)

        # ── Subtitle / metadata ───────────────────────────────────────────────
        meta_lines = []
        if shift_start and shift_end:
            meta_lines.append(
                f"Shift window: {shift_start.strftime('%Y-%m-%d %H:%M UTC')} → "
                f"{shift_end.strftime('%Y-%m-%d %H:%M UTC')}"
            )
        meta_lines.append(f"Generated: {generated_at.strftime('%Y-%m-%d %H:%M UTC')}")

        for line in meta_lines:
            meta_para = doc.add_paragraph()
            meta_run = meta_para.add_run(line)
            meta_run.font.size = Pt(10)
            meta_run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_paragraph_spacing(meta_para, space_before=0, space_after=2)

        # Summary line
        total = sum(len(v) for v in sections.values())
        summary_para = doc.add_paragraph()
        summary_run = summary_para.add_run(
            f"Total items: {total}  |  "
            + "  |  ".join(
                f"{cfg['title'].split('  ')[1]}: {len(sections.get(cfg['key'], []))}"
                for cfg in SECTION_CONFIG
            )
        )
        summary_run.font.size = Pt(9)
        summary_run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
        summary_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(summary_para, space_before=2, space_after=8)

        _add_horizontal_rule(doc)

        # ── Sections ──────────────────────────────────────────────────────────
        for cfg in SECTION_CONFIG:
            section_key = cfg["key"]
            items = sections.get(section_key, [])

            # Section heading
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(cfg["title"])
            heading_run.bold = True
            heading_run.font.size = Pt(14)
            heading_run.font.color.rgb = cfg["heading_color"]
            _set_paragraph_spacing(heading_para, space_before=12, space_after=4)

            if not items:
                # Empty section
                empty_para = doc.add_paragraph()
                empty_run = empty_para.add_run(cfg["empty_msg"])
                empty_run.italic = True
                empty_run.font.size = Pt(10)
                empty_run.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
                _set_paragraph_spacing(empty_para, space_before=0, space_after=4)
            else:
                for item in items:
                    bullet_para = doc.add_paragraph(style="List Bullet")
                    line = _format_event_line(item)
                    bullet_run = bullet_para.add_run(line)
                    bullet_run.font.size = Pt(10)
                    _set_paragraph_spacing(bullet_para, space_before=1, space_after=1)

            _add_horizontal_rule(doc)

        # ── Footer note ───────────────────────────────────────────────────────
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            "This report was auto-generated by the 3HPS Shift Handover System. "
            "Data sourced from: tickets, incidents, chat."
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(0xC0, 0xC0, 0xC0)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_paragraph_spacing(footer_para, space_before=6, space_after=0)

        # ── Save ──────────────────────────────────────────────────────────────
        output_path = str(output_path)
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        doc.save(output_path)

        return output_path

    except Exception as exc:
        raise RuntimeError(
            f"Failed to render shift handover report to '{output_path}': {exc}"
        ) from exc
